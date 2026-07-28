#!/usr/bin/env python3
"""
Script to generate Buku Panduan Peta Tematik Desa Cantik Kabupaten Sidoarjo
in DOCX format, formatted like a User Manual with formal, professional, and comprehensive narrative,
including placeholders for images.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_styled(doc, text, level=1):
    """Add heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run.bold = True
    
    # Adjust spacing
    if level == 1:
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(8)
    else:
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        
    return heading

def add_body_text(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Add body paragraph with consistent font."""
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    # First line indent for standard narrative paragraphs
    if align == WD_ALIGN_PARAGRAPH.JUSTIFY and not bold and not italic:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def add_bullet(doc, text, level=0):
    """Add bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.25 * level)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_numbered(doc, text, level=0):
    """Add numbered list item."""
    p = doc.add_paragraph(style='List Number')
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_image_placeholder(doc, description):
    """Add a visual placeholder for images."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    # Create a visible box representation using a 1x1 table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Set height of the row to give it space
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(r'<w:trHeight {} w:val="2000" w:hRule="atLeast"/>'.format(nsdecls('w')))
    trPr.append(trHeight)
    
    cell = table.cell(0, 0)
    cell.width = Cm(14)
    # Set light gray background
    set_cell_shading(cell, "F2F2F2")
    
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(f"\n[ RUANG UNTUK GAMBAR ]\n{description}\n")
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True
    
    # Add a paragraph after the table
    doc.add_paragraph()

def create_manual():
    doc = Document()

    # =============================================
    # PAGE SETUP
    # =============================================
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(3)

    # =============================================
    # COVER PAGE
    # =============================================
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PETA TEMATIK DESA CANTIK KABUPATEN SIDOARJO\nBuku Petunjuk Penggunaan Aplikasi Berbasis Web")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0, 0, 0)

    for _ in range(12):
        doc.add_paragraph()
        
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = inst.add_run("© Badan Pusat Statistik (BPS) Kabupaten Sidoarjo")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "Hak cipta dilindungi oleh Undang-undang. Buku Petunjuk Penggunaan Web ini dimiliki "
        "oleh Badan Pusat Statistik (BPS) Kabupaten Sidoarjo. Dilarang keras menyalin, "
        "memperbanyak, menerbitkan isi buku ini, sebagian atau seluruhnya, dengan cara apa "
        "pun, baik elektronik maupun mekanis, termasuk tetapi tidak terbatas pada fotokopi atau "
        "sistem penyimpanan data pada komputer."
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    run.italic = True

    doc.add_page_break()

    # =============================================
    # KATA PENGANTAR
    # =============================================
    add_heading_styled(doc, "KATA PENGANTAR", 1)
    
    add_body_text(doc,
        "Puji syukur kami panjatkan ke hadirat Tuhan Yang Maha Esa atas rahmat dan karunia-Nya, "
        "sehingga penyusunan Buku Petunjuk Penggunaan Aplikasi Peta Tematik Desa Cantik Kabupaten "
        "Sidoarjo ini dapat diselesaikan dengan baik. Penyusunan buku panduan ini merupakan "
        "wujud komitmen Badan Pusat Statistik (BPS) Kabupaten Sidoarjo dalam mengoptimalkan "
        "pelayanan informasi geospasial bagi masyarakat, aparatur pemerintah desa, serta "
        "para pemangku kepentingan dalam merumuskan kebijakan berbasis data."
    )
    add_body_text(doc,
        "Penyajian informasi geospasial memegang peranan esensial dalam perencanaan, pelaksanaan, "
        "serta evaluasi program-program pembangunan daerah. Dalam rangka mendukung implementasi "
        "kebijakan Satu Data Indonesia dan percepatan program Kebijakan Satu Peta (One Map Policy), "
        "diperlukan sebuah sistem informasi geospasial yang terintegrasi. Hal ini bertujuan untuk "
        "menyediakan basis data yang konsisten, akurat, dan dapat diakses secara transparan guna "
        "menghindari duplikasi informasi di tingkat wilayah administrasi terkecil."
    )
    add_body_text(doc,
        "Guna merealisasikan sasaran tersebut, BPS Kabupaten Sidoarjo mengembangkan portal Peta "
        "Tematik yang terintegrasi dengan program pembinaan statistik sektoral, yaitu Desa Cantik "
        "(Desa Cinta Statistik). Portal ini dirancang untuk menyajikan representasi visual "
        "berupa sebaran penduduk, struktur demografi, indikator kesejahteraan, hingga inventarisasi "
        "Usaha Mikro, Kecil, dan Menengah (UMKM) secara komprehensif. Kehadiran sistem ini "
        "diharapkan dapat meningkatkan literasi statistik aparatur desa serta mendukung proses "
        "pengambilan keputusan yang presisi dan akuntabel di tingkat daerah."
    )
    add_body_text(doc,
        "Buku panduan ini disusun secara sistematis dan terperinci guna memberikan arahan operasional "
        "dalam pemanfaatan fitur-fitur aplikasi Peta Tematik. Kami berharap dokumen teknis ini "
        "dapat menjadi acuan yang efektif bagi seluruh klasifikasi pengguna, sehingga pemanfaatan "
        "sistem informasi ini dapat berkontribusi maksimal terhadap pembangunan berkelanjutan di "
        "Kabupaten Sidoarjo."
    )
    
    doc.add_paragraph()
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_date.add_run(f"Sidoarjo, {datetime.datetime.now().strftime('%B %Y')}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_sign.add_run("Tim Penyusun")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    
    doc.add_page_break()

    # =============================================
    # DAFTAR ISI
    # =============================================
    add_heading_styled(doc, "DAFTAR ISI", 1)
    
    toc_items = [
        ("KATA PENGANTAR", 1),
        ("DAFTAR ISI", 1),
        ("PENDAHULUAN", 1),
        ("A. Latar Belakang Program Desa Cantik", 2),
        ("B. Tujuan Sistem Peta Tematik", 2),
        ("PANDUAN PENGGUNAAN", 1),
        ("A. Cara Mengakses Situs (Login Umum)", 2),
        ("B. Penggunaan Menu Utama", 2),
        ("    1. Halaman Peta Statistik (Beranda)", 3),
        ("    2. Halaman Peta Tematik Desa", 3),
        ("    3. Menu Pusat Bantuan Terpadu", 3),
        ("C. Fungsionalitas Penampil Peta (Map Viewer)", 2),
        ("    1. Navigasi Peta Interaktif", 3),
        ("    2. Penggunaan Instrumen Peta (Map Widgets)", 3),
        ("    3. Visualisasi Data Tingkat Rukun Tetangga (RT)", 3),
        ("    4. Implementasi Fitur AI Insight", 3),
        ("D. Panduan Dasbor Administrator", 2),
        ("    1. Prosedur Autentikasi (Login)", 3),
        ("    2. Tata Letak Antarmuka Dasbor Admin", 3),
        ("    3. Pemutakhiran Data dan Manajemen Insight", 3),
    ]

    for title_text, level in toc_items:
        p = doc.add_paragraph()
        if level == 1:
            run = p.add_run(f"{title_text}")
            run.bold = True
        elif level == 2:
            run = p.add_run(f"{title_text}")
            p.paragraph_format.left_indent = Cm(0.5)
        else:
            run = p.add_run(f"{title_text}")
            p.paragraph_format.left_indent = Cm(1.0)
            
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # =============================================
    # PENDAHULUAN
    # =============================================
    add_heading_styled(doc, "PENDAHULUAN", 1)

    add_heading_styled(doc, "A. Latar Belakang Program Desa Cantik", 2)
    add_body_text(doc,
        "Pembangunan nasional yang berkelanjutan mengandalkan ketersediaan data yang valid, akurat, "
        "dan representatif. Sejalan dengan implementasi Peraturan Presiden Nomor 39 Tahun 2019 "
        "tentang Satu Data Indonesia, pemerintah menetapkan kerangka kerja tata kelola data yang "
        "mendukung pembagian informasi secara efektif antar instansi pemerintah. Dalam konteks ini, "
        "desa—sebagai entitas pemerintahan terkecil yang bersinggungan langsung dengan dinamika "
        "masyarakat—dituntut untuk tidak hanya bertindak sebagai objek pengumpulan data, melainkan "
        "berperan aktif sebagai subjek pengelola data yang mencerminkan realitas sosial ekonomi wilayahnya."
    )
    add_body_text(doc,
        "Merespons kebutuhan tata kelola data tersebut, Badan Pusat Statistik (BPS) Republik Indonesia "
        "meluncurkan program pembinaan statistik sektoral yang bertajuk Desa Cantik (Desa Cinta Statistik). "
        "Program strategis ini diinisiasi dengan sasaran utama untuk meningkatkan kapasitas, literasi, "
        "serta partisipasi aparatur pemerintah desa dalam penyelenggaraan statistik. Pemahaman mendalam "
        "terhadap data kewilayahan memungkinkan pemerintah desa untuk mengidentifikasi permasalahan secara "
        "objektif dan menyusun perencanaan pembangunan yang tepat sasaran (evidence-based policy), "
        "seperti alokasi bantuan sosial yang akurat, pembangunan infrastruktur yang terukur, hingga "
        "strategi pemberdayaan ekonomi masyarakat setempat."
    )
    add_body_text(doc,
        "Di wilayah Kabupaten Sidoarjo, BPS mengimplementasikan semangat Desa Cantik melalui pengembangan "
        "sistem informasi yang inovatif dan terintegrasi. Kabupaten Sidoarjo memiliki lanskap demografi "
        "dan ekonomi yang heterogen, meliputi kawasan perindustrian padat karya, pusat perdagangan, "
        "hingga sentra produksi agrikultur dan perikanan. Keberagaman karakteristik wilayah ini menuntut "
        "suatu metode penyajian informasi statistik yang mampu memvisualisasikan keunggulan spesifik dari "
        "masing-masing daerah secara presisi."
    )
    add_body_text(doc,
        "Metodologi penyajian data secara konvensional, yang umumnya berbasis pada format tabulasi dan "
        "dokumen cetak statis, dinilai memiliki keterbatasan dalam mengakomodasi dinamika informasi masa kini. "
        "Format tersebut cenderung membutuhkan waktu interpretasi yang lebih lama dan kurang efektif "
        "dalam merepresentasikan korelasi spasial antarwilayah. Mengatasi tantangan tersebut, BPS Kabupaten "
        "Sidoarjo mengembangkan Aplikasi Web Peta Tematik Desa Cantik, sebuah platform Sistem Informasi "
        "Geografis (SIG) interaktif yang memproyeksikan indikator statistik—mulai dari kepadatan demografi, "
        "kelayakan perumahan, hingga pemetaan sektor riil—langsung ke dalam koordinat geografis hingga level "
        "Rukun Tetangga (RT)."
    )

    add_heading_styled(doc, "B. Tujuan Sistem Peta Tematik", 2)
    add_body_text(doc,
        "Tujuan fundamental dari pengembangan platform Peta Tematik ini adalah memperkuat literasi data "
        "publik dan instansi pemerintahan melalui penyajian visual yang komprehensif. Dengan "
        "menstransformasikan basis data statistik menjadi peta tematik bergradasi warna (choropleth), "
        "diagram struktural, dan analisis naratif yang didukung oleh Kecerdasan Buatan (AI), aplikasi "
        "ini menjembatani kesenjangan interpretasi data yang kerap dialami oleh pengguna non-statistik."
    )
    add_body_text(doc,
        "Lebih jauh, aplikasi ini difungsikan sebagai sistem pendukung keputusan (decision support system) "
        "esensial bagi para pengambil kebijakan. Sebagai contoh, aparatur pemerintah dapat secara instan "
        "memvisualisasikan area dengan disparitas rasio jenis kelamin yang tinggi atau mengidentifikasi "
        "konsentrasi lokasi Usaha Mikro, Kecil, dan Menengah (UMKM) dalam suatu wilayah. Akses informasi "
        "real-time yang terpusat ini merupakan wujud dedikasi BPS Kabupaten Sidoarjo dalam mendukung tata "
        "kelola pemerintahan yang transparan, akuntabel, dan didasari pada pembuktian empiris."
    )

    doc.add_page_break()

    # =============================================
    # PANDUAN PENGGUNAAN
    # =============================================
    add_heading_styled(doc, "PANDUAN PENGGUNAAN", 1)
    
    add_body_text(doc,
        "Bagian ini mendeskripsikan prosedur teknis pengoperasian fitur-fitur yang tersedia pada "
        "Aplikasi Peta Tematik Desa Cantik bagi pengguna umum. Sistem ini telah didesain dengan "
        "antarmuka yang ergonomis untuk memastikan navigasi dan perolehan informasi dapat "
        "dilaksanakan secara efisien."
    )

    add_heading_styled(doc, "A. Cara Mengakses Situs (Login Umum)", 2)
    add_body_text(doc,
        "Aplikasi ini dikategorikan sebagai platform informasi publik yang dapat diakses melalui "
        "jaringan internet secara global tanpa memerlukan registrasi awal bagi pengguna umum. "
        "Prosedur akses adalah sebagai berikut:"
    )
    add_numbered(doc, "Persiapan Perangkat Lunak: Buka peramban web (web browser) pilihan Anda. Untuk menjamin seluruh modul pemetaan dan grafik divisualisasikan dengan spesifikasi optimal, direkomendasikan menggunakan peramban mutakhir seperti Google Chrome, Mozilla Firefox, Microsoft Edge, atau Safari.")
    add_numbered(doc, "Alokasi URL: Ketikkan alamat URL (Uniform Resource Locator) resmi dari Peta Tematik Desa Cantik Kabupaten Sidoarjo pada bilah alamat peramban.")
    add_numbered(doc, "Akses Halaman Beranda: Tekan tombol Enter. Sistem akan memuat antarmuka pengguna secara otomatis dan mengarahkan Anda ke Halaman Beranda yang memuat Peta Statistik Kabupaten Sidoarjo. Tidak diperlukan proses autentikasi pada tahap ini.")

    add_heading_styled(doc, "B. Penggunaan Menu Utama", 2)
    add_body_text(doc,
        "Aplikasi dilengkapi dengan bilah navigasi (navigation bar) statis di bagian atas antarmuka, "
        "berfungsi sebagai panel kendali utama untuk mengakses berbagai modul dalam sistem."
    )
    
    add_heading_styled(doc, "1. Halaman Peta Statistik (Beranda)", 3)
    add_body_text(doc,
        "Halaman ini menyajikan visualisasi makro wilayah Kabupaten Sidoarjo berdasarkan segmentasi "
        "administratif tingkat Kecamatan. Modul ini menggunakan metode pemetaan choropleth, yakni teknik "
        "pewarnaan area berdasarkan klasifikasi data statistik."
    )
    add_body_text(doc,
        "Pengguna memiliki kapabilitas untuk beralih antara berbagai variabel pengukuran, "
        "seperti metrik 'Kepadatan Penduduk' maupun 'Rasio Jenis Kelamin'. Halaman ini dilengkapi pula "
        "dengan fungsi pencarian dinamis (search box) yang akan secara otomatis menyesuaikan titik "
        "koordinat peta (zoom to bounding box) pada kecamatan yang dituju, serta menampilkan panel ringkasan "
        "demografi agregat di sisi layar."
    )
    add_image_placeholder(doc, "Screenshot tampilan Beranda Peta Statistik Kabupaten Sidoarjo")

    add_heading_styled(doc, "2. Halaman Peta Tematik Desa", 3)
    add_body_text(doc,
        "Dengan memilih menu 'Peta Tematik', sistem akan merender ulang batas spasial menuju tingkat "
        "pemerintahan yang lebih mikro, yakni level Desa/Kelurahan. Tujuan operasional dari modul ini "
        "adalah menyoroti wilayah-wilayah yang telah tergabung sebagai wilayah percontohan Desa Cantik."
    )
    add_body_text(doc,
        "Desa-desa yang telah menginventarisasi potensi sektoralnya akan diwakili dengan pewarnaan yang "
        "kontras (emas/kuning). Melalui interaksi klik pada poligon desa yang menyala tersebut, "
        "sistem akan menampilkan kartu informasi ringkas yang memuat profil potensi desa dan menyediakan "
        "tautan langsung menuju Halaman Detail Desa terkait."
    )
    add_image_placeholder(doc, "Screenshot tampilan Halaman Peta Tematik tingkat Desa/Kelurahan")

    add_heading_styled(doc, "3. Menu Pusat Bantuan Terpadu", 3)
    add_body_text(doc,
        "Menu ini dilambangkan dengan ikon tanda tanya (?). Modul ini didedikasikan sebagai pusat rujukan "
        "operasional (help desk), yang memuat daftar Pertanyaan yang Sering Diajukan (Frequently Asked Questions), "
        "penjelasan metodologi pewarnaan peta, serta panduan teknis pelaporan melalui narahubung (hotline) "
        "resmi BPS Kabupaten Sidoarjo."
    )

    add_heading_styled(doc, "C. Fungsionalitas Penampil Peta (Map Viewer)", 2)
    add_body_text(doc,
        "Intrumen visualisasi utama pada sistem ini adalah kanvas pemetaan spasial berbasis vektor yang "
        "mendukung interaksi penuh dengan pengguna."
    )
    
    add_heading_styled(doc, "1. Navigasi Peta Interaktif", 3)
    add_body_text(doc,
        "Pengguna dapat melaksanakan perintah pergeseran bidang peta (panning) dengan cara mengklik dan "
        "menahan tombol utama tetikus sembari menggeser kursor. Fungsi perbesaran atau pengecilan skala "
        "pandang (zooming) dapat dioperasikan menggunakan tuas gulir (scroll wheel). Apabila pengguna "
        "mengarahkan kursor (hover) di atas suatu poligon wilayah, sistem secara seketika mengeksekusi "
        "permintaan informasi (tooltip) yang menampilkan identitas wilayah dan nilai metrik utama secara "
        "real-time."
    )
    add_image_placeholder(doc, "Screenshot fitur interaksi hover/tooltip pada area peta")

    add_heading_styled(doc, "2. Penggunaan Instrumen Peta (Map Widgets)", 3)
    add_body_text(doc,
        "Sistem ini menyertakan sejumlah kontrol operasional (widget) yang terintegrasi di sisi layar:"
    )
    add_bullet(doc, "Pemilih Lapisan Dasar (Basemap Selector): Terletak di sudut kanan atas antarmuka, memberikan opsi bagi pengguna untuk beralih representasi latar peta, yang meliputi Peta Standar (vektor jalan dan label), Citra Satelit (tangkapan visual dari udara), dan Peta Topografi.")
    add_bullet(doc, "Kontrol Skala (Zoom Controls): Tombol operasional berlambang plus (+) dan minus (-) yang dirancang untuk memanipulasi resolusi pandang peta secara presisi.")
    add_bullet(doc, "Panel Legenda (Legend): Terletak berdekatan dengan alat pengontrol skala. Panel ini bertugas sebagai indikator rentang nilai statistik yang dikorelasikan dengan gradasi warna pada peta choropleth.")
    add_bullet(doc, "Panel Saring Data (Filter Data): Instrumen berbentuk kotak dialog (panel) yang memfasilitasi pengguna dalam melakukan penapisan variabel spesifik yang ingin divisualisasikan pada layar (misal: kategori usia tertentu atau jenis lapangan usaha).")

    add_heading_styled(doc, "3. Visualisasi Data Tingkat Rukun Tetangga (RT)", 3)
    add_body_text(doc,
        "Pada Halaman Detail Desa, fungsionalitas peta diperluas hingga mencapai batas spasial level "
        "Rukun Tetangga (RT). Saat pengguna berinteraksi dengan salah satu poligon RT, modul grafik "
        "akan dieksekusi."
    )
    add_body_text(doc,
        "Visualisasi kuantitatif ini diwujudkan melalui serangkaian diagram interaktif, yang mencakup "
        "Piramida Penduduk (histogram demografi bilateral), Grafik Lingkaran (Pie Chart) untuk representasi "
        "distribusi material hunian, dan penghitung angka statis beranimasi (CountUp) yang merepresentasikan "
        "agregat data makro seperti total tenaga kerja pada area tersebut."
    )
    add_image_placeholder(doc, "Screenshot tampilan grafik dan piramida penduduk pada Halaman Detail Desa")

    add_heading_styled(doc, "4. Implementasi Fitur AI Insight", 3)
    add_body_text(doc,
        "Inovasi analisis tingkat lanjut dihadirkan melalui modul AI Insight, yang direpresentasikan "
        "oleh ikon bintang/robot (✨) pada bagian sudut bawah aplikasi. Fitur ini dirancang untuk "
        "menterjemahkan kompleksitas data kuantitatif menjadi paparan teks deskriptif yang komprehensif."
    )
    add_body_text(doc,
        "Terdapat dua kategori insight yang disajikan: 'Insight Manual' yang merepresentasikan evaluasi "
        "kualitatif resmi dari aparatur desa terkait, dan 'Insight AI' yang diolah secara komputasional "
        "oleh pemodelan bahasa berskala besar (Large Language Model). Sistem akan menganalisis tren dan "
        "korelasi data di area terpilih, kemudian menyusun satu paragraf analisis otomatis yang objektif."
    )
    add_image_placeholder(doc, "Screenshot tampilan panel AI Insight yang sedang menampilkan analisis teks")

    doc.add_page_break()

    # =============================================
    # PANDUAN ADMIN
    # =============================================
    add_heading_styled(doc, "D. Panduan Dasbor Administrator", 2)
    
    add_body_text(doc,
        "Dalam rangka memastikan integritas dan akurasi data yang dipublikasikan, sistem menyediakan "
        "lingkungan tertutup berupa Dasbor Administrator. Ruang kerja digital ini diperuntukkan "
        "eksklusif bagi administrator yang memiliki otorisasi, meliputi pihak BPS (Admin Pusat) "
        "dan perangkat pemerintah desa (Admin Desa)."
    )

    add_heading_styled(doc, "1. Prosedur Autentikasi (Login)", 3)
    add_body_text(doc,
        "Sistem menerapkan mekanisme autentikasi terpusat (single-gateway authentication) dengan prosedur "
        "sebagai berikut:"
    )
    add_numbered(doc, "Inisiasi Akses: Klik tombol 'Masuk Admin' yang terletak secara persisten di sebelah kanan bilah navigasi atas.")
    add_numbered(doc, "Verifikasi Kredensial: Pada halaman autentikasi, inputkan Nama Pengguna (Username) dan Kata Sandi (Password) yang sah sesuai dengan otorisasi yang diterbitkan oleh BPS Kabupaten Sidoarjo.")
    add_numbered(doc, "Eksekusi Otorisasi: Tekan tombol eksekusi login. Berdasarkan kredensial yang divalidasi, sistem akan melakukan pengalihan (routing) spesifik: administrator BPS diarahkan ke konsol Admin Pusat, sedangkan perangkat desa diarahkan ke panel kerja instansi desanya masing-masing.")
    add_image_placeholder(doc, "Screenshot halaman formulir Login Admin")

    add_heading_styled(doc, "2. Tata Letak Antarmuka Dasbor Admin", 3)
    add_body_text(doc,
        "Setelah autentikasi berhasil, administrator akan dipaparkan pada antarmuka manajemen berbasis "
        "tabulasi logis untuk efisiensi alur kerja."
    )
    add_bullet(doc, "Konsol Admin Pusat: Administrator BPS diberikan kapabilitas administratif superior, termasuk pembuatan akun otentikasi (provisioning) bagi entitas desa baru, serta penentuan 'Tema Potensi' untuk tiap entitas (misal: Sektor Kependudukan, Ekonomi UMKM, atau Agrikultur). Konfigurasi pada parameter ini akan memicu perubahan struktur pemetaan di aplikasi publik secara real-time.")
    add_bullet(doc, "Panel Admin Desa: Lingkungan kerja bagi aparatur desa didesain dengan tingkat kustomisasi dinamis yang menyesuaikan dengan Tema Potensi desa bersangkutan. Opsi manajemen (tab) yang disajikan akan dieliminasi apabila tidak berkesesuaian dengan fokus pemetaan data desa tersebut, sehingga meminimalisasi redudansi operasional.")
    add_image_placeholder(doc, "Screenshot tampilan antarmuka Dasbor Admin Desa")

    add_heading_styled(doc, "3. Pemutakhiran Data dan Manajemen Insight", 3)
    add_body_text(doc,
        "Fungsi operasional pokok bagi Administrator Desa terklasifikasi dalam dua ranah pengelolaan:"
    )
    add_bullet(doc, "Pemeliharaan Rekod Data (Tabel): Tabulator manajemen data menyediakan antarmuka terstruktur serupa lembar kerja elektronik (spreadsheet). Administrator berkewajiban melakukan penambahan, penyuntingan, atau penghapusan rekod statistik sektoral (seperti kuantitas hunian, unit usaha, maupun parameter spasial lainnya). Algoritma penyimpanan akan memproses transisi data secara sinkron dengan antarmuka publik.")
    add_bullet(doc, "Penyusunan Manajemen Insight: Guna melengkapi data kuantitatif dengan konteks empiris kewilayahan, administrator disediakan instrumen input teks naratif. Catatan strategis, evaluasi peristiwa, atau justifikasi atas suatu lonjakan statistik dapat diinput pada modul ini. Data kualitatif tersebut akan diekspos kepada publik melalui antarmuka 'Insight Manual', melengkapi analisis yang dibangkitkan oleh sistem kecerdasan buatan.")
    add_image_placeholder(doc, "Screenshot tampilan pengelolaan tabel data dan pengisian teks Insight")

    # =============================================
    # SAVE
    # =============================================
    output_path = r"D:\Documents\Internship\BPS Sidoarjo\desa-cantik\Buku_Petunjuk_Penggunaan_Web_Peta_Tematik_v3.docx"
    doc.save(output_path)
    print(f"Buku panduan profesional dengan placeholder gambar berhasil dibuat: {output_path}")
    return output_path

if __name__ == "__main__":
    create_manual()
